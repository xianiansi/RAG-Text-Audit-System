import streamlit as st
from docx import Document
import re
import json
import dashscope
from dashscope import Generation
import os,dotenv


# 函数：根据段落内容来识别段落类型
def classify_paragraph(doc, para):
    text = para.text.strip()
    if not text:
        return None
    if doc.is_appendix:
        return "appendix_table"
    if text == "附件":
        doc.is_appendix = True
        return "appendix_table"
    if text in doc.title_parts:
        return "title"
    elif re.match(r"稽审任字[^\s]+号", text):
        return "doc_number"
    elif re.match(r"^.{3,15}营业部$", text):
        return "sub_title"
    elif re.match(r"^原负责人.{2,4}同志离任审计报告$", text):
        return "sub_title"
    elif re.match(r"^负责人.{2,4}同志离岗审计报告$", text):
        return "sub_title"
    elif text == "董事长、总审计师、党委组织部-人力资源部：":
        return "report_subject"
    elif re.match(r"^[一二三四五六七八九十]+、", text):
        return "level1_heading"
    elif re.match(r"^（[一二三四五六七八九十]+）", text):
        return "level2_heading"
    elif re.match(r"^[1-9][0-9]*、", text):
        return "level3_heading"
    else:
        return "main_text"


# 函数：提取段落格式
def extract_format(para):
    para_format = {}

    # 默认字体信息
    font_name_CN = None
    font_name_num = None

    # 遍历段落的每个 run
    for run in para.runs:
        if run.text.strip():  # 检查 run 中是否有非空白字符
            # 判断字符是否为中文
            if any('\u4e00' <= char <= '\u9fff' for char in run.text):
                font_name_CN = run.font.name
            # 判断字符是否为英文或数字
            elif any('A' <= char <= 'Z' or 'a' <= char <= 'z' or '0' <= char <= '9' for char in run.text):
                font_name_num = run.font.name

    # 如果未找到字体，尝试使用段落样式的字体
    if not font_name_CN:
        font_name_CN = para.style.font.name if para.style.font.name else '宋体'  # 使用段落样式的字体作为备选
    if not font_name_num:
        font_name_num = para.style.font.name if para.style.font.name else 'Times New Roman'

    # 对于所有字符的字体样式进行综合处理
    para_format['font_name_CN'] = font_name_CN
    para_format['font_name_num'] = font_name_num
    para_format['font_size'] = para.runs[0].font.size if para.runs else para.style.font.size
    para_format['line_spacing'] = para.paragraph_format.line_spacing
    para_format['first_line_indent'] = para.paragraph_format.first_line_indent
    para_format['alignment'] = para.alignment

    return para_format


# 类：文档分类和格式提取
class DocClassification:
    def __init__(self, doc_path):
        self.doc = Document(doc_path)  # 加载模板文档
        self.section_format = {}  # 初始化分类字典
        self.title_parts = ["国泰君安证券股份有限公司", "审计报告书"]  # 标题检查列表
        self.is_appendix = False  # 初始化标志位：标识是否进入附表部分

    # 遍历文档中的每个段落，分类并填入记录类别的格式
    def classify_and_extract(self):
        for para in self.doc.paragraphs:
            para_type = classify_paragraph(self, para)
            if para_type:  # 跳过空段落
                format_info = extract_format(para)
                self.section_format[para_type] = format_info

        # 打印分类识别结果（可根据需要保存到文件）
        output_path = 'document_sections.json'
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.section_format, f, ensure_ascii=False, indent=4)

        st.write(f"文档段落格式已保存到 {output_path}")

    def get_full_text(self):
        """获取文档的完整文本内容"""
        full_text = "\n".join([para.text for para in self.doc.paragraphs if para.text.strip()])
        return full_text


# 函数：对比两个文档的格式
def compare_formats(current_format, template_format):
    mismatched_paragraphs = []

    for para_type, format_info in current_format.items():
        if para_type in template_format:
            expected_format = template_format[para_type]

            # 比较当前文档的段落格式与模板格式
            for key in expected_format:
                if key in format_info and format_info[key] != expected_format[key]:
                    mismatched_paragraphs.append((para_type, key, format_info[key], expected_format[key]))

    return mismatched_paragraphs


# Streamlit 应用设置
st.title('文件审计对比')

# 上传第一个模板文件
uploaded_template = st.file_uploader("请上传审计模板文件", type=["docx"])

if uploaded_template is not None:
    # 处理模板文件并保存其格式
    doc_classifier = DocClassification(uploaded_template)
    doc_classifier.classify_and_extract()

    st.write("模板文件已处理并保存格式信息。")

    # 上传要对比的第二个文件
    uploaded_doc = st.file_uploader("请上传要对比的审计文件", type=["docx"])

    if uploaded_doc is not None:
        # 处理要对比的文件并保存其格式
        doc_to_compare = DocClassification(uploaded_doc)
        doc_to_compare.classify_and_extract()

        # 读取模板的格式信息
        with open('document_sections.json', 'r', encoding='utf-8') as f:
            template_format = json.load(f)

        # 对比格式并显示不符合要求的段落
        mismatched_paragraphs = compare_formats(doc_to_compare.section_format, template_format)

        if mismatched_paragraphs:
            st.write("以下段落格式不符合模板要求：")
            for para_type, key, current, expected in mismatched_paragraphs:
                st.write(f"段落类型: {para_type}")
                st.write(f"- 属性: {key}, 当前: {current}, 期望: {expected}")
        else:
            st.write("所有段落格式均符合模板要求。")

            # 添加一个文本输入框
            user_input = st.text_input("请输入其他审计要求:")

            # 添加“开始审计”按钮
            if st.button("开始审计"):
                # 调用外部服务进行生成
                dotenv.load_dotenv()
                dashscope.api_key = os.getenv("dashscope_api_key")
                prompt = ("你是一个金融领域审计专家，你需要对以下文件内容进行审核，审核内容为："
                          + user_input + " 需要审计的文件内容为：" + doc_to_compare.get_full_text())

                # 调用大模型生成结果
                responses = Generation.call(model=Generation.Models.qwen_turbo, prompt=prompt)

                # 获取生成的文本
                generated_text = responses.output['text']

                # 在页面上显示生成的文本
                st.subheader("审计结果：")
                st.text_area("生成的审计报告", value=generated_text, height=300)

                # 将生成的文本保存为Markdown文件
                output_file_path = '生成的文档.md'
                with open(output_file_path, 'w', encoding='utf-8') as f:
                    f.write(generated_text)

                st.write(f"生成的文本已保存为 {output_file_path}")



